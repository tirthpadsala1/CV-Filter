from PyPDF2 import PdfReader 
import os
import sys
import gc
from docx import Document

from utils.logging import logging
from utils.exception import CustomException

def textExtract_pdf(path):
        
        try:
            text = ""

            with open(path, 'rb') as pdf_file:
                reader = PdfReader(pdf_file)

                if reader.is_encrypted:
                    logging.info(f"removed encrypted file:{path}")
                    return text
            
                for page in reader.pages:
                    text += page.extract_text()

            del reader; gc.collect()
            logging.info(f"exracted text from :{path} , text:{text}")
            return text.strip()

        except Exception as e:
            logging.info(f"function name: [ textExtract_pdf() ] exception occured:{e} ")
            raise CustomException(e,sys)
        
def textExtract_docx(path):
        try:
            doc = Document(path)
            lines = []
            for p in doc.paragraphs:
                if p.text.strip():
                    lines.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        lines.append(" | ".join(row_text))

            del doc; gc.collect()
            logging.info(f"exracted text from :{path} , text:{lines}")
            return "\n".join(lines)

        except Exception as e:
            if any(keyword in str(e) for keyword in  ['encrypt', 'password', 'protected'] ):
                logging.info(f"removed encrypted file:{path}")  
                return ""     

            logging.info(f"function name: [ textExtract_docx() ] exception occured:{e} ")
            raise CustomException(e,sys)
        