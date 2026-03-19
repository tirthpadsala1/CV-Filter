from utils.var import content

from utils.logging import logging
from utils.exception import CustomException

from PyPDF2 import PdfReader
from docx import Document
from pathlib import Path
from itertools import zip_longest
import os
import shutil
import gc
import sys
import re

class CVClassifier:

    def __init__(self , downloadFolder):
        self.downloadFolder = downloadFolder
        self.destinationPath = content["CVFolder"]

    def textExtract_pdf(self , path):
        
        try:
            text = ""

            with open(path, 'rb') as pdf_file:
                reader = PdfReader(pdf_file)

                if reader.is_encrypted:
                    logging.info(f"removed encrypted file:{path}")
                    return text
            
                for page in reader.pages:
                    text += page.extract_text()

            del reader; gc.collect()
            logging.info(f"exracted text from :{path} , text:{text}")
            return text.strip()

        except Exception as e:
            logging.info(f"function name: [ textExtract_pdf() ] exception occured:{e} ")
            raise CustomException(e,sys)
        
    def textExtract_docx(self , path):
        try:
            doc = Document(path)
            lines = []
            for p in doc.paragraphs:
                if p.text.strip():
                    lines.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        lines.append(" | ".join(row_text))

            del doc; gc.collect()
            logging.info(f"exracted text from :{path} , text:{lines}")
            return "\n".join(lines)

        except Exception as e:
            if any(keyword in str(e) for keyword in  ['encrypt', 'password', 'protected'] ):
                logging.info(f"removed encrypted file:{path}")  
                return ""     

            logging.info(f"function name: [ textExtract_docx() ] exception occured:{e} ")
            raise CustomException(e,sys)
        
    def classifier(self , filename:str ,  text:str):

        try:


            keywords = {
                'high': ['resume','cv','curriculum vitae','curriculum','work experience','professional experience','employment history','education','qualifications','skills','expertise','summary','objective','profile','career','references','certifications','projects','achievements','career summary','professional summary','technical skills','core competencies','key skills','areas of expertise','experience','academic background','education background','training','internship','internships','publications','research','conferences','awards','honors','responsibilities','employment','job history','work history','career history','certifications & courses','professional affiliations'],
                
                'medium': ['name','address','phone','email','contact','university','college','degree','bachelor','master','phd','diploma','certificate','language','programming','technical','soft skills','github','linkedin','portfolio','coursework','gpa','cgpa','grade','projects','intern','trainee','developer','engineer','analyst','manager','consultant','researcher','technologies','tools','frameworks','stack','methodologies','achieved','led','designed','implemented','built','deployed','maintained','responsible for','volunteer','extracurricular','hackathon','competition','open source','certified','credential','publication','journal','thesis']
            }

            filenames = [
                    r'.*resume.*',r'.*cv.*',r'.*curriculum.*vitae.*',r'.*curriculum.*',r'.*profile.*',r'.*resume.*portfolio.*',r'.*job.*application.*',r'.*employment.*history.*',r'.*work.*experience.*',r'.*professional.*summary.*',r'.*career.*objective.*',r'^[a-z]+_[a-z]+_(resume|cv)',r'^(resume|cv)_.*',r'^[a-z]+\s[a-z]+\s(resume|cv)', r'^(resume|cv)\s.*',r'^[a-z]+\s[a-z]+_?(resume|cv)',r'^[a-z]+\s[a-z]+_?cv_?.*',r'^[a-z]+\s[a-z]+_?resume_?.*',r'.*(resume|cv)\s*20\d{2}.*',r'.*(resume|cv)\s*\d{4}.*',r'.*(resume|cv)\s*\d{2}\.\d{2}\.\d{4}.*', r'.*(resume|cv)\s*v\d+.*',r'.*(resume|cv)\s*version\d+.*',r'.*(resume|cv)\s*final.*',r'.*(resume|cv)\s*updated.*',r'.*technical.*resume.*',r'.*academic.*cv.*',r'.*executive.*resume.*',r'.*student.*resume.*',r'.*graduate.*cv.*',r'.*my\s*(resume|cv).*',r'.*personal\s*(resume|cv).*',r'.*professional\s*(resume|cv).*',r'.*latest\s*(resume|cv).*',r'.*current\s*(resume|cv).*',r'^[a-z]+\s*_\s*[a-z]+\s*_\s*(resume|cv)\s*',r'^[a-z]+\s+[a-z]+\s*-\s*(resume|cv)\s*',
                    r'^[a-z]+\s*\.\s*[a-z]+\s*\.\s*(resume|cv)\s*',
                    r'.*(resume|cv)\s*for\s+.*',
                    r'.*(resume|cv)\s*application.*',
                    r'.*(resume|cv)\s*submission.*',
                    r'.*résumé.*',
                    r'.*res.*',
                    r'.*bio.*data.*',
                    r'.*resume.*(?:\.\w+)?',
                    r'.*cv.*(?:\.\w+)?'
                    ]
            for pattern in filenames:
                if re.match(pattern , filename):
                    logging.info(f"potential CV , name:{filename} , pattern:{pattern} ")
                    return True
                
            text = re.sub(r"[^a-z0-9\s]", " ", text)
            words = text.split()

            bigrams = [f"{words[i]} {words[i+1]}" for i in range(len(words)-1)]

            score = 0
            matched = set()

            for w,bg in zip_longest(words , bigrams , fillvalue="xx"):
                if w in keywords["high"] and w not in matched:
                    score += 3
                    matched.add(w)
                if w in keywords["medium"] and w not in matched:
                    score += 1
                    matched.add(w)

                if bg in keywords["high"] and bg not in matched:
                    score += 3
                    matched.add(bg)

                if bg in keywords["medium"] and bg not in matched:
                    score += 1
                    matched.add(bg)

            if score >= 15:
                    logging.info(f"potential CV , name:{filename} , score:{score} ")
                    return True

            elif score >= 8:
                    logging.info(f"potential CV , name:{filename} , score:{score} ")
                    return True

            else:
                logging.info(f"not a CV , name:{filename} , score:{score} ")
                return False
            

            
            
        except Exception as e:
            logging.info(f"function name: [ classifier() ] exception occured:{e} ")
            raise CustomException(e,sys)
        
    def DirectoryLoop(self):

        try:
            
            folder = Path(self.downloadFolder)
            nonCV = []
            CV = []

            validExtension = {".pdf" , ".docx"}
            for path in folder.rglob("*"):
                if path.suffix.lower() in validExtension:
                    fileName = os.path.basename(path)
                    if path.suffix.lower() == ".pdf":
                        text = self.textExtract_pdf(path).lower()
                    else:
                        text = self.textExtract_docx(path).lower()

                    Result  = self.classifier(fileName , text)

                    if Result == False:
                        nonCV.append(path)

                    elif text == "":
                        nonCV.append(path)
                        logging.info(f"encrypted file:{fileName} , path:{path}")
                    else:
                        CV.append(path)
                        

            actionPerformed = {
                "movedFiles" : [],
                "deletedFiles" : []            
                }

            for i in CV:
                
                shutil.move(i , self.destinationPath)
                logging.info(f'''passed file:{fileName} , from:{path} , to:{content["CVFolder"]}''')
                actionPerformed["movedFiles"].append(i)

            for i in nonCV:

                os.remove(i)
                logging.info(f"deleted file:{fileName} , path:{path}")
                actionPerformed["deletedFiles"].append(i)
            
            return actionPerformed


            


        except Exception as e:
            logging.info(f"function name: [ DirectoryLoop() ] exception occured:{e} ")
            raise CustomException(e,sys)
        
    